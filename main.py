from flask import Flask, request, send_file, jsonify
from docx import Document
import os
import tempfile
import logging
from pathlib import Path

# ---------------- LOGGING ----------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s"
)

app = Flask(__name__)

# ---------------- CONFIG ----------------
TEMPLATE_PATH = Path(__file__).parent / "templates" / "FrontPage_Template.docx"

# ---------------- HELPERS ----------------
def replace_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, val)

    if new_text != full_text:
        # remove old runs
        for run in paragraph.runs[::-1]:
            run._element.getparent().remove(run._element)
        paragraph.add_run(new_text)

def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, replacements)

def replace_everywhere(doc, replacements):
    # normal paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)

    # tables
    for table in doc.tables:
        replace_in_table(table, replacements)

    # headers & footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, replacements)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, replacements)

# ---------------- ROUTES ----------------
@app.route("/", methods=["GET"])
def health():
    logging.info("❤️ Health check hit")
    return "OK", 200

@app.route("/generate", methods=["POST"])
def generate_docx():
    try:
        logging.info("📥 /generate called")

        data = request.get_json(force=True)
        logging.info(f"📦 Payload received: {data}")

        if not TEMPLATE_PATH.exists():
            logging.error(f"❌ Template missing at {TEMPLATE_PATH}")
            return jsonify(error="Template DOCX not found"), 500

        replacements = {
            "{{TEST_NAME}}": data.get("test", ""),
            "{{CLASS}}": data.get("class", ""),
            "{{PHASE}}": data.get("phase", ""),
            "{{SET}}": data.get("set", ""),
            "{{DATE}}": data.get("date", "")
        }

        logging.info("📄 Loading template DOCX")
        doc = Document(TEMPLATE_PATH)

        logging.info("🔁 Replacing placeholders")
        replace_everywhere(doc, replacements)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp_path = tmp.name
        tmp.close()

        doc.save(tmp_path)
        logging.info(f"✅ DOCX generated at {tmp_path}")

        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=f"FrontPage_{data.get('set','SET')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logging.exception("❌ DOCX generation failed")
        return jsonify(error=str(e)), 500
